#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历风险评估脚本 v2 - 改进版
针对PDF提取文本质量差的情况优化
"""

import sys
import json
import re
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── 风险维度定义 ──────────────────────────────────────────────────────────────

RISK_LEVELS = ['低', '中', '高']
RISK_COLORS = {'低': 'C6EFCE', '中': 'FFEB9C', '高': 'FFC7CE'}
RISK_FONTS = {'低': '006100', '中': '9C6500', '高': '9C0006'}


def parse_resume(text, input_name=None):
    """解析简历文本，提取关键字段"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    clean_lines = [re.sub(r'[\~\`\(\)\;\•\–\-]', '', l).strip() for l in lines]
    
    result = {
        'name': input_name or '',  # 优先使用输入的名字
        'education': [],
        'work_experience': [],
        'skills': [],
        'raw_text': text,
        'clean_text': ' '.join(clean_lines)
    }
    
    # 如果没有输入名字，则尝试从文本提取
    if not input_name:
        # 提取姓名 - 在整个文本中搜索（PDF文字乱码，姓名位置不固定）
        name_patterns = [
            r'姓\s*名[：:\s]*([\u4e00-\u9fa5]{2,4})',
            r'([\u4e00-\u9fa5]{2,3})\s{1,3}\d{3}[-\d]{7,}',
            r'^([\u4e00-\u9fa5]{2,3})\s+(?:个人简历|简历|求职)',
            r'([\u4e00-\u9fa5]{2,4})\s{0,3}[_~\d](?:\s|$)',
        ]
        found_name = None
        for line in lines:
            for p in name_patterns:
                m = re.search(p, line)
                if m:
                    found_name = m.group(1)
                    break
            if found_name:
                break
        
        if not found_name:
            # 备选：在clean_text中搜索第一个2-4字中文名
            m = re.search(r'([\u4e00-\u9fa5]{2,4})', result['clean_text'][:200])
            if m:
                found_name = m.group(1)
        
        result['name'] = found_name or '未知'
    else:
        result['name'] = input_name
    
    # 提取学历信息
    edu_pattern = r'(\d{4})\s*[-–至至]\s*(\d{4}|\d{2})[年.]?\s*[\u4e00-\u9fa5]{2,10}(?:大学|学院|学校|研究生|硕士|博士|本科|大专|中专)'
    for line in clean_lines:
        m = re.search(edu_pattern, line)
        if m:
            result['education'].append(line[:80])
    
    # 提取工作经历（按行）
    work_section = False
    work_keywords = ['工作经历', '任职', '职位', '公司']
    for line in clean_lines:
        if any(k in line[:8] for k in work_keywords) and len(line) < 20:
            work_section = True
            continue
        if work_section and len(line) > 10:
            result['work_experience'].append(line[:100])
        elif work_section and not line:
            work_section = False
    
    # 提取技能
    skill_section = False
    skill_keywords = ['个人技能', '专业技能', '技术专长', '擅长']
    for line in clean_lines:
        if any(k in line[:8] for k in skill_keywords) and len(line) < 15:
            skill_section = True
            continue
        if skill_section and line:
            result['skills'].append(line[:80])
        elif skill_section and len(line) < 5:
            skill_section = False
    
    return result


def analyze_risk(resume_data):
    """分析简历风险，返回风险评估结果"""
    risks = {
        'education': {'level': '低', 'flags': []},
        'work_experience': {'level': '低', 'flags': []},
        'timeline': {'level': '低', 'flags': []},
        'salary': {'level': '低', 'flags': []},
        'skill_exaggeration': {'level': '低', 'flags': []},
        'vague_language': {'level': '低', 'flags': []}
    }
    
    text = resume_data['clean_text']
    work_exp = resume_data.get('work_experience', [])
    edu = resume_data.get('education', [])
    name = resume_data.get('name', '')
    
    # ── 时间线分析（改进版，减少误报）──
    year_pattern = r'(?:20)1[5-9]|20[2-3]\d'  # 只匹配合理的年份
    years = re.findall(year_pattern, text)
    years = sorted(set([int(y) for y in years]))
    
    if len(years) >= 2:
        year_counts = {}
        for y in years:
            year_counts[y] = text.count(str(y))
        
        # 只在年数 >= 4 且有明显重叠时报警
        if len(years) >= 4:
            max_count = max(year_counts.values())
            min_count = min(year_counts.values())
            if max_count > 3 and max_count / min_count > 2:
                risks['timeline']['level'] = '中'
                risks['timeline']['flags'].append(f'工作时间段可能存在重叠，年份出现频率异常')
    
    # ── 职位/公司模糊分析 ─
    vague_patterns = [
        (r'某[知名大型热门]\w*', '包含"某知名/大型"等模糊公司描述'),
        (r'\w{2,4}互联网\w{0,4}', '公司描述过于泛化'),
        (r'重要项目|核心产品|关键任务', '关键职责描述模糊'),
    ]
    for p, msg in vague_patterns:
        if re.search(p, text):
            risks['vague_language']['level'] = '中'
            if msg not in risks['vague_language']['flags']:
                risks['vague_language']['flags'].append(msg)
            break
    
    # ── 数字完美性分析 ─
    perfect_numbers = re.findall(r'(?:100%|99%|98%|97%|96%|95%|增长了\s*\d+%|提升了\s*\d+%)', text)
    if len(perfect_numbers) > 3:
        risks['skill_exaggeration']['level'] = '中'
        risks['skill_exaggeration']['flags'].append(f"简历含{len(perfect_numbers)}处完美数据，建议核实")
    
    # ── 技能夸大分析 ─
    exaggeration_phrases = [
        "精通", "熟练掌握", "深度理解", "全面掌握", "专家级", "顶级",
    ]
    exp_count = 0
    for phrase in exaggeration_phrases[:4]:  # 只检查前4个
        exp_count += len(re.findall(phrase, text))
    
    if exp_count >= 6:
        risks['skill_exaggeration']['level'] = '高'
        risks['skill_exaggeration']['flags'].append(f'简历中{exp_count}处使用夸大技能描述')
    elif exp_count >= 4:
        risks['skill_exaggeration']['level'] = '中'
        risks['skill_exaggeration']['flags'].append(f'多处使用"精通"等夸大用词（{exp_count}处）')
    
    # ── 工作经历风险 ─
    if len(work_exp) == 0:
        risks['work_experience']['level'] = '中'
        risks['work_experience']['flags'].append('未识别到工作经历描述')
    elif len(work_exp) < 3 and len(text) > 1500:
        risks['work_experience']['level'] = '中'
        risks['work_experience']['flags'].append('工作经历描述偏少，与简历篇幅不符')
    
    # ── 学历风险 ─
    if len(edu) == 0:
        risks['education']['level'] = '低'  # 不强制报警，因为PDF可能提取不全
        risks['education']['flags'].append('未识别到明确学历信息（建议人工核实）')
    
    # ── 计算综合风险（修复bug）──
    level_to_num = {'低': 0, '中': 1, '高': 2}
    total_score = sum(level_to_num[r['level']] for r in risks.values())
    flagged_count = sum(1 for r in risks.values() if r['flags'])
    
    # 综合风险：考虑flagged数量和最高风险维度
    max_risk_num = max(level_to_num[r['level']] for r in risks.values())
    
    overall = '低'
    if max_risk_num >= 2:  # 有高风险维度
        overall = '高'
    elif max_risk_num == 1 or flagged_count >= 3:  # 有中风险或3个以上flag
        overall = '中'
    
    return {
        'name': name or '未知',
        'risks': risks,
        'overall': overall,
        'summary': '；'.join([f"[{k}]{v['level']}" for k, v in risks.items() if v['flags']]) or '无明显风险点'
    }


def create_excel_report(results, output_path):
    """生成Excel风险评估报告"""
    wb = Workbook()
    
    # ── Sheet 1: 风险总览 ─────────────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = '风险总览'
    
    headers = ['序号', '姓名', '整体风险', '学历', '经历', '时间线', '薪资', '技能夸大', '表述模糊', '风险点摘要']
    header_fill = PatternFill('solid', fgColor='4472C4')
    header_font = Font(bold=True, color='FFFFFF')
    
    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    for idx, r in enumerate(results, 1):
        row = [
            idx,
            r['name'],
            r['overall'],
            r['risks']['education']['level'],
            r['risks']['work_experience']['level'],
            r['risks']['timeline']['level'],
            r['risks']['salary']['level'],
            r['risks']['skill_exaggeration']['level'],
            r['risks']['vague_language']['level'],
            r['summary']
        ]
        for col, val in enumerate(row, 1):
            cell = ws1.cell(row=idx+1, column=col, value=val)
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            if col in [3, 4, 5, 6, 7, 8, 9] and val in RISK_COLORS:
                cell.fill = PatternFill('solid', fgColor=RISK_COLORS[val])
                cell.font = Font(color=RISK_FONTS[val])
    
    widths = [6, 12, 8, 8, 8, 8, 8, 8, 8, 50]
    for i, w in enumerate(widths, 1):
        ws1.column_dimensions[get_column_letter(i)].width = w
    ws1.row_dimensions[1].height = 25
    
    # ── Sheet 2: 详细分析 ─────────────────────────────────────────────────────
    ws2 = wb.create_sheet('详细分析')
    
    headers2 = ['序号', '姓名', '风险维度', '风险等级', '具体风险点']
    for col, h in enumerate(headers2, 1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    
    row_num = 2
    for idx, r in enumerate(results, 1):
        for dim, data in r['risks'].items():
            if data['flags']:
                for flag in data['flags']:
                    ws2.cell(row=row_num, column=1, value=idx)
                    ws2.cell(row=row_num, column=2, value=r['name'])
                    ws2.cell(row=row_num, column=3, value=dim)
                    ws2.cell(row=row_num, column=4, value=data['level'])
                    ws2.cell(row=row_num, column=5, value=flag)
                    
                    cell = ws2.cell(row=row_num, column=4)
                    if data['level'] in RISK_COLORS:
                        cell.fill = PatternFill('solid', fgColor=RISK_COLORS[data['level']])
                        cell.font = Font(color=RISK_FONTS[data['level']])
                    
                    row_num += 1
    
    for i, w in enumerate([6, 12, 15, 8, 60], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    
    # ── Sheet 3: 汇总统计 ─────────────────────────────────────────────────────
    ws3 = wb.create_sheet('汇总统计')
    
    ws3['A1'] = '简历风险评估汇总'
    ws3['A1'].font = Font(bold=True, size=16)
    
    ws3['A3'] = '风险等级分布'
    ws3['A3'].font = Font(bold=True, size=12)
    
    risk_counts = {'高': 0, '中': 0, '低': 0}
    for r in results:
        risk_counts[r['overall']] += 1
    
    ws3['A5'] = '风险等级'
    ws3['B5'] = '人数'
    ws3['C5'] = '占比'
    for cell in [ws3['A5'], ws3['B5'], ws3['C5']]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill('solid', fgColor='D9E1F2')
    
    for i, (level, count) in enumerate(risk_counts.items(), 6):
        ws3.cell(row=i, column=1, value=level)
        ws3.cell(row=i, column=2, value=count)
        ws3.cell(row=i, column=3, value=f'=B{i}/SUM(B$6:B$8)')
        ws3.cell(row=i, column=3).number_format = '0.0%'
        
        cell = ws3.cell(row=i, column=1)
        cell.fill = PatternFill('solid', fgColor=RISK_COLORS[level])
        cell.font = Font(color=RISK_FONTS[level])
    
    # 高风险名单
    ws3['A11'] = '高风险候选人'
    ws3['A11'].font = Font(bold=True, size=12)
    ws3['A12'] = '序号'
    ws3['B12'] = '姓名'
    ws3['C12'] = '主要风险点'
    for cell in [ws3['A12'], ws3['B12'], ws3['C12']]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill('solid', fgColor='D9E1F2')
    
    hi_risk = [r for r in results if r['overall'] == '高']
    for i, r in enumerate(hi_risk, 13):
        ws3.cell(row=i, column=1, value=i-12)
        ws3.cell(row=i, column=2, value=r['name'])
        ws3.cell(row=i, column=3, value=r['summary'][:60])
    
    # 中风险名单
    mid_risk = [r for r in results if r['overall'] == '中']
    if mid_risk:
        start_row = 13 + len(hi_risk) + 2
        ws3.cell(row=start_row, column=1, value='中风险候选人')
        ws3.cell(row=start_row, column=1).font = Font(bold=True, size=12)
        ws3.cell(row=start_row+1, column=1, value='序号')
        ws3.cell(row=start_row+1, column=2, value='姓名')
        ws3.cell(row=start_row+1, column=3, value='主要风险点')
        for cell in [ws3.cell(row=start_row+1, column=1), ws3.cell(row=start_row+1, column=2), ws3.cell(row=start_row+1, column=3)]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill('solid', fgColor='D9E1F2')
        for i, r in enumerate(mid_risk, start_row+2):
            ws3.cell(row=i, column=1, value=i-start_row-1)
            ws3.cell(row=i, column=2, value=r['name'])
            ws3.cell(row=i, column=3, value=r['summary'][:60])
    
    for i, w in enumerate([8, 15, 60], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    
    wb.save(output_path)
    return output_path


def process_folder(folder_path):
    """处理文件夹中的所有简历文件"""
    import os
    results = []
    
    # 支持的文件扩展名
    supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls']
    
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext not in supported_extensions:
                continue
            
            file_path = os.path.join(root, file)
            print(f"处理: {file}")
            
            try:
                text = ""
                if ext == '.pdf':
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                elif ext in ['.docx', '.doc']:
                    from docx import Document
                    doc = Document(file_path)
                    text = "\n".join([p.text for p in doc.paragraphs])
                elif ext in ['.xlsx', '.xls']:
                    import openpyxl
                    wb = openpyxl.load_workbook(file_path, read_only=True)
                    for sheet in wb:
                        for row in sheet.iter_rows(values_only=True):
                            for cell in row:
                                if cell:
                                    text += str(cell) + " "
                            text += "\n"
                
                # 使用文件名作为候选姓名
                name = os.path.splitext(file)[0]
                parsed = parse_resume(text, name)
                analyzed = analyze_risk(parsed)
                results.append(analyzed)
                print(f"  → 完成: {analyzed['name']} - {analyzed['overall']}风险")
                
            except Exception as e:
                print(f"  → 处理失败: {str(e)}")
                continue
    
    return results


def main():
    import argparse
    parser = argparse.ArgumentParser(description='简历风险批量评估工具')
    parser.add_argument('--file', help='单个简历文件路径')
    parser.add_argument('--folder', help='简历文件夹路径')
    args = parser.parse_args()
    
    results = []
    
    try:
        if args.folder:
            # 处理文件夹
            results = process_folder(args.folder)
        elif args.file:
            # 处理单个文件
            import os
            ext = os.path.splitext(args.file)[1].lower()
            text = ""
            if ext == '.pdf':
                import pdfplumber
                with pdfplumber.open(args.file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            elif ext in ['.docx', '.doc']:
                from docx import Document
                doc = Document(args.file)
                text = "\n".join([p.text for p in doc.paragraphs])
            
            name = os.path.splitext(os.path.basename(args.file))[0]
            parsed = parse_resume(text, name)
            analyzed = analyze_risk(parsed)
            results.append(analyzed)
        else:
            # 原来的JSON输入方式
            input_data = sys.stdin.buffer.read()
            try:
                resumes = json.loads(input_data.decode('utf-8'))
                if not isinstance(resumes, list):
                    resumes = [resumes]
                for resume in resumes:
                    text = resume.get('text', resume.get('content', '')) if isinstance(resume, dict) else str(resume)
                    name = resume.get('name') if isinstance(resume, dict) else None
                    parsed = parse_resume(text, name)
                    analyzed = analyze_risk(parsed)
                    results.append(analyzed)
            except json.JSONDecodeError as e:
                print(json.dumps({'success': False, 'error': f'Invalid JSON: {e}'}))
                sys.exit(1)
        
        output_path = 'resume_risk_report.xlsx'
        create_excel_report(results, output_path)
        
        print(json.dumps({
            'success': True,
            'output': output_path,
            'count': len(results),
            'high_risk': len([r for r in results if r['overall'] == '高']),
            'medium_risk': len([r for r in results if r['overall'] == '中']),
            'low_risk': len([r for r in results if r['overall'] == '低']),
            'details': [{'name': r['name'], 'overall': r['overall'], 'summary': r['summary']} for r in results]
        }, ensure_ascii=False))
        
    except Exception as e:
        print(json.dumps({'success': False, 'error': str(e)}))
        sys.exit(1)


if __name__ == '__main__':
    main()
